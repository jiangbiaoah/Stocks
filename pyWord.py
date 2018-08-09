# -*- coding: utf-8 -*-
import json
import datetime
from docx import Document
from MyStocks import remaindecimal
from docx.shared import RGBColor, Pt


class MyWord():
    current = datetime.datetime.now()
    doc = None
    pars = None
    result = None

    def __init__(self):
        self.doc = Document('./example.docx')

    def generate(self, text):
        self.result = json.loads(text)
        self.pars = self.doc.paragraphs
        self.__change_title()
        self.__change_szzs()
        self.__change_szcz()
        self.__change_hszs()
        self.__change_table_1()
        self.__change_table_2()

        filepath = './files/' + 'A股公司及同业公司股价表现-%s%s%s%s%s%s.docx' % (self.current.year, self.current.month, self.current.day, self.current.hour, self.current.minute, self.current.second)
        self.doc.save(filepath)
        print("文档输出完成")
        return filepath

    def __change_title(self):
        # for par in pars:
        #     print(par.text)
        title = self.pars[0]
        title.text = ''
        run = title.add_run('%s年%s月%s日' % (self.current.year, self.current.month, self.current.day))
        run.font.size = Pt(16)

    def __change_szzs(self):    # 上证指数 p = 3
        par = self.pars[3]
        par.text = ''
        run = par.add_run('上证指数收  ')
        run.font.size = Pt(14)
        run.font.bold = True

        run = par.add_run('%s， %s， %s%%' % (remaindecimal(self.result['szzs']['nowpri']), remaindecimal(self.result['szzs']['increase'], symbol=1), remaindecimal(self.result['szzs']['increPer'], symbol=1)))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.__get_font_color(float(self.result['szzs']['increase']))

    def __change_szcz(self):    # 深证成指 p = 4
        par = self.pars[4]
        par.text = ''
        run = par.add_run('深证成指收  ')
        run.font.size = Pt(14)
        run.font.bold = True

        run = par.add_run('%s， %s， %s%%' % (remaindecimal(self.result['szcz']['nowpri']), remaindecimal(self.result['szcz']['increase'], symbol=1), remaindecimal(self.result['szcz']['increPer'], symbol=1)))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.__get_font_color(float(self.result['szcz']['increase']))

    def __change_hszs(self):    # 恒生指数 p = 6
        par = self.pars[6]
        par.text = ''
        run = par.add_run('恒生指数收  ')
        run.font.size = Pt(14)
        run.font.bold = True

        run = par.add_run('%s， %s， %s%%' % (remaindecimal(self.result['hszs']['nowpri']), remaindecimal(self.result['hszs']['increase'], symbol=1), remaindecimal(self.result['hszs']['increPer'], symbol=1)))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.__get_font_color(float(self.result['hszs']['increase']))

    def __change_table_1(self):
        table_1 = self.doc.tables[0]

        stocks_a = ['sh601992', 'sz000401', 'sz000856', 'sh600585', 'sz000002']
        var = -1
        for row in table_1.rows:
            cell = row.cells
            if var == -1:
                var = 0
                continue
            cell[0].text = ''
            cell[1].text = ''   # CNY
            cell[2].text = ''
            cell[3].text = ''
            cell[4].text = ''
            cell[5].text = ''
            cell[6].text = ''
            cell[7].text = ''

            run = cell[0].paragraphs[0].add_run(self.result['company'][stocks_a[var]]['stockName'])
            run.font.size = Pt(12)

            run = cell[1].paragraphs[0].add_run('CNY')
            run.font.size = Pt(12)

            run = cell[2].paragraphs[0].add_run(remaindecimal(self.result['company'][stocks_a[var]]['currentPrice']))
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_a[var]]['changeAmount']))

            run = cell[3].paragraphs[0].add_run(remaindecimal(self.result['company'][stocks_a[var]]['changeAmount'], symbol=1))
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_a[var]]['changeAmount']))

            run = cell[4].paragraphs[0].add_run(remaindecimal(self.result['company'][stocks_a[var]]['priceChangeRatio'], symbol=1))
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_a[var]]['changeAmount']))

            run = cell[5].paragraphs[0].add_run(remaindecimal(str(float(self.result['company'][stocks_a[var]]['amount'])/100000000)))
            run.font.size = Pt(12)

            var = var + 1

    def __change_table_2(self):
        table_2 = self.doc.tables[1]

        stocks_hk = ['H02009', 'H00914', 'H03323', 'H01313', 'H00688']
        var = -1
        for row in table_2.rows:
            cell = row.cells
            if var == -1:
                var = 0
                continue
            # for stock in stocks_hk:
            cell[0].text = ''
            cell[1].text = ''
            cell[2].text = ''
            cell[3].text = ''
            cell[4].text = ''
            cell[5].text = ''

            run = cell[0].paragraphs[0].add_run(self.result['company'][stocks_hk[var]]['stockName'])
            run.font.size = Pt(12)

            run = cell[1].paragraphs[0].add_run('HKD')
            run.font.size = Pt(12)

            run = cell[2].paragraphs[0].add_run(remaindecimal(self.result['company'][stocks_hk[var]]['currentPrice']))
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_hk[var]]['changeAmount']))

            run = cell[3].paragraphs[0].add_run(remaindecimal(self.result['company'][stocks_hk[var]]['changeAmount'], symbol=1))
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_hk[var]]['changeAmount']))

            run = cell[4].paragraphs[0].add_run(self.result['company'][stocks_hk[var]]['priceChangeRatio'])
            run.font.size = Pt(12)
            run.font.color.rgb = self.__get_font_color(float(self.result['company'][stocks_hk[var]]['changeAmount']))

            amount = self.result['company'][stocks_hk[var]]['amount']
            # pattern = re.compile(r'\d+')
            # code = pattern.findall(amount)[0]
            number = None
            if amount.find(u'万') != -1:
                index = amount.find(u'万')
                number = amount[0:index]
                number = remaindecimal(str(float(number)/10000))
            else:
                index = amount.find(u'亿')
                number = amount[0:index]
                number = remaindecimal(number)

            run = cell[5].paragraphs[0].add_run(number)
            run.font.size = Pt(12)

            var = var + 1

    def __get_font_color(self, number):
        if number < 0:
            return RGBColor(0, 255, 0)  # 绿色
        elif number > 0:
            return RGBColor(255, 0, 0)  # 红色
        else:
            return RGBColor(0, 0, 0)  # 黑色

