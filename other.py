# -*- coding: utf-8 -*-

import datetime
import json
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from MyStocks import remaindecimal
'''
pip install python-docx
Document -> Paragraph -> Run
'''


class MyWord():

    def __init__(self):
        print("")

    def export(self, text):
        result = json.loads(text)
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
        doc.styles['Normal'].font.size = Pt(12)

        current = datetime.datetime.now()

        p00 = doc.add_paragraph()
        p00.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p00.add_run("%s年%s月%s日" % (current.year, current.month, current.day) + '\n' + "公司及同业公司股价表现")
        run.font.size = Pt(16)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

        p1 = doc.add_paragraph()
        run = p1.add_run("（一）A股")
        run.font.size = Pt(16)

        p11 = doc.add_paragraph()
        run = p11.add_run("上证指数收  ")
        run.font.size = Pt(14)
        run.font.bold = True
        run = p11.add_run("%s, %s, %s%%" % (remaindecimal(result['szzs']['nowpri']), remaindecimal(result['szzs']['increase']), remaindecimal(result['szzs']['increPer'])))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = getFontColor(float(result['szzs']['increase']))

        p12 = doc.add_paragraph()
        run = p12.add_run("深证成指收  ")
        run.font.size = Pt(14)
        run.font.bold = True
        run = p12.add_run("%s, %s, %s%%" % (remaindecimal(result['szcz']['nowpri']), remaindecimal(result['szcz']['increase']), remaindecimal(result['szcz']['increPer'])))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = getFontColor(float(result['szcz']['increase']))

        table1 = doc.add_table(rows=6, cols=8, style="Light List Accent 2")
        cells1_0 = table1.rows[0].cells
        cells1_0[0].text = '公司'
        cells1_0[1].text = '货币'
        cells1_0[2].text = '今日收市价'
        cells1_0[3].text = '相比前一交易日'
        cells1_0[4].text = '涨跌幅度'
        cells1_0[5].text = '成交金额'
        cells1_0[6].text = '换手率'
        cells1_0[7].text = '总市值（亿元）'
        # stocks_a = ['H02009', 'H00914', 'H03323', 'H01313', 'H00688']
        stocks_a = ['金隅集团', ['冀东水泥'], ['冀东装备'], ['海螺水泥'], ['万科A']]
        i = 1
        for stock in stocks_a:
            table1.rows[i].cells[0].text = stock
            table1.rows[i].cells[1].text = 'CNY'
            table1.rows[i].cells[2].text = ''
            table1.rows[i].cells[3].text = ''
            table1.rows[i].cells[4].text = ''
            table1.rows[i].cells[5].text = ''
            table1.rows[i].cells[6].text = ''
            table1.rows[i].cells[7].text = ''
            i = i + 1

        p2 = doc.add_paragraph()
        run = p2.add_run("（二）H股")
        run.font.size = Pt(16)

        p21 = doc.add_paragraph()
        run = p21.add_run("恒生指数收  ")
        run.font.size = Pt(14)
        run.font.bold = True
        run = p21.add_run("%s, %s, %s%%" % (remaindecimal(result['hszs']['nowpri']), remaindecimal(result['hszs']['increase']), remaindecimal(result['hszs']['increPer'])))
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = getFontColor(float(result['hszs']['increase']))

        table2 = doc.add_table(rows=6, cols=6, style="Light List Accent 2")
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells2_0 = table2.rows[0].cells
        cells2_0[0].text = '公司'
        cells2_0[1].text = '货币'
        cells2_0[2].text = '今日收市价'
        cells2_0[3].text = '相比前一交易日'
        cells2_0[4].text = '涨跌幅度'
        cells2_0[5].text = '成交金额'

        stocks_hk = ['H02009', 'H00914', 'H03323', 'H01313', 'H00688']
        j = 1
        for stock in stocks_hk:
            table2.rows[j].cells[0].text = result['company'][stock]['stockName']
            table2.rows[j].cells[1].text = 'HKD'
            # table2.rows[j].cells[2].text = remaindecimal(result['company'][stock]['currentPrice'])
            # table2.rows[j].cells[3].text = remaindecimal(result['company'][stock]['changeAmount'])
            # table2.rows[j].cells[4].text = result['company'][stock]['priceChangeRatio']
            table2.rows[j].cells[5].text = result['company'][stock]['amount']

            run = table2.cell(j, 2).paragraphs[0].add_run(remaindecimal(result['company'][stock]['currentPrice']))
            run.font.color.rgb = getFontColor(float(result['company'][stock]['changeAmount']))
            run = table2.cell(j, 3).paragraphs[0].add_run(remaindecimal(result['company'][stock]['changeAmount']))
            run.font.color.rgb = getFontColor(float(result['company'][stock]['changeAmount']))
            run = table2.cell(j, 4).paragraphs[0].add_run(result['company'][stock]['priceChangeRatio'])
            run.font.color.rgb = getFontColor(float(result['company'][stock]['changeAmount']))

            j = j + 1

        doc.save('A股公司及同业公司股价表现-%s%s%s%s%s%s.docx' % (current.year, current.month, current.day, current.hour, current.minute, current.second))
        print("文档输出完成")


def getFontColor(number):
    if number < 0:
        return RGBColor(0, 255, 0)  # 绿色
    elif number > 0:
        return RGBColor(255, 0, 0)  # 红色
    else:
        return RGBColor(0, 0, 0)  # 黑色

